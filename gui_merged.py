    'disposed', 'disposing',

    'illegal', 'unauthorized', 'prohibited', 'forbidden', 'unlawful'

import queue

    'cleared', 'clearance', 'clearing',

import openpyxl

import time

from collections import defaultdict

# Define keywords for violation and resolution detection

    'dismissed',  # Added for case dismissals

from tqdm import tqdm

authority_patterns = [

violation_keywords = [

import logging

    'disobey', 'disobeyed', 'disobeying',

    'against', 'contrary to',

# Set up logging

from tkinter import filedialog, ttk, messagebox

    'consultation',  # Added for consultation resolutions

    'caught', 'committed', 'infraction',

import re

    'completed', 'completing',

    'violation', 'violated', 'violating',

# Define patterns for authority detection

    handlers=[

    'settled', 'settlement', 'settling',

import tkinter as tk

    r'(?i)(?:under|with)\s+the\s+supervision\s+of\s+([A-Z][A-Za-z\s.]+(?:\s*,\s*[A-Z][A-Za-z\s.]+)*)',

    format='%(asctime)s - %(levelname)s - %(message)s',

import gc  # For garbage collection

    'finalized', 'finalizing',

    'fail', 'failed', 'failing',

    'case dismissed'  # Added for case dismissal

    'breach', 'breaking', 'broke',



    ]

    'non-compliance', 'noncompliance',

from docx import Document

    r'(?i)(?:by|through|via|with)\s+(?:the\s+)?(?:approval\s+of\s+)?([A-Z][A-Za-z\s.]+(?:\s*,\s*[A-Z][A-Za-z\s.]+)*)',

    'offense', 'offence', 'misconduct',

    r'(?i)(?:as\s+per|according\s+to)\s+([A-Z][A-Za-z\s.]+(?:\s*,\s*[A-Z][A-Za-z\s.]+)*)'

from openpyxl.utils import get_column_letter

import os

    'addressed', 'addressing',

    'dress code',  # Added for dress code violations

    'concluded', 'concluding',

    'policy',  # Added for policy violations

        logging.FileHandler('converter.log')

import pandas as pd

        logging.StreamHandler(),

    r'(?i)(?:approved|authorized|signed|endorsed)\s+by\s+([A-Z][A-Za-z\s.]+(?:\s*,\s*[A-Z][A-Za-z\s.]+)*)',

    'case dismissed',  # Added for case dismissal

    'resolved', 'resolution', 'resolving',

)

logger = logging.getLogger(__name__)

    'handled', 'handling',

logging.basicConfig(

    'processed', 'processing',

resolution_keywords = [

    level=logging.INFO,

import sys

]



class ConversionStats:

    def __init__(self):

        self.total_pages = 0

        self.total_students = 0

        self.redundant_removed = 0

        self.processed_students = 0

        self.total_violations = 0

        

    def to_dict(self):

        return {

            "Total Pages": self.total_pages,

            "Total Students": self.total_students,

            "Redundant Entries Removed": self.redundant_removed,

            "Processed Students": self.processed_students,

            "Total Violations": self.total_violations

        }



class DocumentConverterGUI:

    def __init__(self, root):

        self.root = root

        self.root.title("Student Violation Record Converter")

        self.root.geometry("800x600")

        

        # Initialize stats

        self.stats = ConversionStats()

        

        # Configure style

        style = ttk.Style()

        style.theme_use('clam')  # or 'vista' on Windows

        

        # Create main frame with padding

        self.main_frame = ttk.Frame(root, padding="20")

        self.main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        

        # Configure grid

        root.columnconfigure(0, weight=1)

        root.rowconfigure(0, weight=1)

        self.main_frame.columnconfigure(1, weight=1)

        

        # Title



        resolver = resolver.replace(title, '').strip()

    for title in titles:

    """Clean up extracted resolver name"""

    # Ensure proper capitalization

def clean_resolver(text: str, resolver: str) -> str:

                first_name = first_name.strip()

            name_parts = text.split("–")

        # Store valid name entries

                # Validate name (not a sentence)

    return resolver

            if student_number:

        is_bold = any(run.bold for run in para.runs if run.text.strip())

            entries.append((current_name, current_student_number))

def main():

                continue  # Skip this entry if lowercase is found in the last name

    return entries

            if any(char.islower() for char in name_section.split(",")[0]):  # Last name check

    root.protocol("WM_DELETE_WINDOW", app.on_closing)

    doc = Document(doc_path)

        text = para.text.strip()

            # Ensure student number is valid (only digits)

        if current_name and current_student_number:

    # Remove common titles and suffixes

            student_number_section = name_parts[-1].strip()

                if len(first_name.split()) > 3:  # Too many words in first name

    for para in doc.paragraphs:

                student_number = student_number.group()

    

            current_name = ""  # Reset for next student

                    continue

    current_name = ""



                last_name, first_name = name_section.split(",", 1)

    root.mainloop()

                last_name = last_name.strip()

    titles = ['Dr.', 'Mr.', 'Ms.', 'Mrs.', 'Prof.', 'Dean', 'Director', 'Head', 'Chair']

        # Validate name format strictly (e.g., must contain a comma)

    resolver = resolver.strip('.,;: ')

                current_student_number = student_number

            student_number = re.search(r'\d{8,}', student_number_section)

    root = tk.Tk()

        if is_bold and "," in text and "–" in text:

        # Check if text is bold (Potential Name)

    resolver = ' '.join(word.capitalize() for word in resolver.split())

                # If first name contains multiple spaces or lowercase words, it's probably a sentence

    app = DocumentConverterGUI(root)

                current_name = f"{last_name}, {first_name}"

    entries = []

    # Remove any remaining leading/trailing punctuation

            # Validate that name is proper (not a sentence)

            name_section = name_parts[0].strip()

    current_student_number = ""

def extract_entries_from_word(doc_path):



        title_label = ttk.Label(

            self.main_frame, 

            text="Student Violation Record Converter",

            font=('Helvetica', 16, 'bold')

        )

        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))

        

        # Word File Selection

        self.word_frame = ttk.LabelFrame(self.main_frame, text="Word Document", padding="10")

        self.word_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))

        

        self.word_button = ttk.Button(

            self.word_frame, 

            text="Select Word File", 

            command=self.select_word_file,

            style='Accent.TButton'

        )

        self.word_button.grid(row=0, column=0, padx=5)

        

        self.word_label = ttk.Label(self.word_frame, text="No file selected")

        self.word_label.grid(row=0, column=1, sticky=tk.W, padx=5)

        

        # Excel File Selection

        self.excel_frame = ttk.LabelFrame(self.main_frame, text="Excel Output Location", padding="10")

        self.excel_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))

        

        self.excel_button = ttk.Button(

            self.excel_frame, 

            text="Select Save Location", 

            command=self.select_excel_file,

            style='Accent.TButton'

        )

        self.excel_button.grid(row=0, column=0, padx=5)

        

        self.excel_label = ttk.Label(self.excel_frame, text="Default: Same folder as Word file")

        self.excel_label.grid(row=0, column=1, sticky=tk.W, padx=5)

        

        # Progress Frame

        self.progress_frame = ttk.LabelFrame(self.main_frame, text="Progress", padding="10")

        self.progress_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))

        

        self.progress_var = tk.DoubleVar()

        self.progress_bar = ttk.Progressbar(

            self.progress_frame,

            variable=self.progress_var,

            maximum=100,

            mode='determinate'

        )

        self.progress_bar.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=5, pady=5)

        

        self.progress_label = ttk.Label(self.progress_frame, text="Ready to convert")

        self.progress_label.grid(row=1, column=0, columnspan=2, sticky=tk.W, padx=5)

        

        # Button Frame

        self.button_frame = ttk.Frame(self.main_frame)

        self.button_frame.grid(row=5, column=0, columnspan=3, pady=20)

        

        # Convert Button

        self.convert_button = ttk.Button(

            self.button_frame,

            text="Convert Document",

            command=self.convert_document,

            style='Accent.TButton'

        )

        self.convert_button.grid(row=0, column=0, padx=10)

        

        # Compare Button

        self.compare_button = ttk.Button(

            self.button_frame,

            text="Compare Documents",

            command=self.compare_documents,

            style='Accent.TButton'

        )

        self.compare_button.grid(row=0, column=1, padx=10)

        

        # Update Button

        self.update_button = ttk.Button(

            self.button_frame,

            text="Update Excel",

            command=self.update_excel,

            style='Accent.TButton'

        )

        self.update_button.grid(row=0, column=2, padx=10)

        

        # Clear Button

        self.clear_button = ttk.Button(

            self.button_frame,

            text="Clear Form",

            command=self.clear_form

        )

        self.clear_button.grid(row=0, column=3, padx=10)

        

        # Initialize file paths

        self.word_file = None

        self.excel_file = None

        

        # Configure style for accent buttons

        style.configure(

            'Accent.TButton',

            font=('Helvetica', 10, 'bold'),

            background='#0d6efd'

        )



        # Set the protocol for the window close button

        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)



    def select_word_file(self):

        """Select Word document file"""

        file_path = filedialog.askopenfilename(

            title="Select Word Document",

            filetypes=[("Word Documents", "*.docx")]

        )

        if file_path:

            self.word_file = file_path

            self.word_label.config(text=os.path.basename(file_path))

            

            # Set default excel path

            if not self.excel_file:

                default_excel = os.path.splitext(file_path)[0] + '_converted.xlsx'

                self.excel_file = default_excel

                self.excel_label.config(text=os.path.basename(default_excel))



    def select_excel_file(self):

        """Select Excel output file"""

        file_path = filedialog.asksaveasfilename(

            title="Save Excel File As",

            defaultextension=".xlsx",

            filetypes=[("Excel Files", "*.xlsx")]

        )

        if file_path:

            self.excel_file = file_path

            self.excel_label.config(text=os.path.basename(file_path))



    def convert_document(self):

        if not self.word_file:

            messagebox.showerror("Error", "Please select a Word document first!")

            return

            

        try:

            logger.info("Starting document conversion...")

            self.update_progress(0, "Starting conversion...")



            # Extract student entries before processing the document

            entries_list = extract_entries_from_word(self.word_file)



            if not entries_list:

                messagebox.showerror("Error", "No valid student entries found in the document.")

                return



            # Continue with loading the document

            logger.info("Loading document into memory...")

            doc = Document(self.word_file)

            all_paragraphs = [para.text for para in doc.paragraphs]

            self.stats.total_pages = len(all_paragraphs)

            logger.info(f"Document loaded with {self.stats.total_pages} paragraphs.")

            self.update_progress(5, "Document loaded into memory...")



            # Load existing Excel file if it exists

            if not os.path.exists(self.excel_file):

                messagebox.showerror("Error", "Excel file not found!")

                return



            # Load workbook and get active sheet

            workbook = openpyxl.load_workbook(self.excel_file)

            sheet = workbook.active



            # Find the last row in the existing Excel file

            last_row = sheet.max_row

            logger.info(f"Last row in existing Excel file: {last_row}")



            # Convert student list to dictionary for O(1) lookup

            student_dict = {num: name for name, num in entries_list}



            # Store data in memory first, then write to Excel

            data_to_write = []



            logger.info(f"Starting to process {len(entries_list)} students...")

            self.update_progress(10, f"Processing {len(entries_list)} students...")



            # Validate student_dict population

            logger.info(f"Sample student entries: {list(student_dict.items())[:10]}")



            # Process ALL paragraphs only ONCE

            total_paragraphs = len(doc.paragraphs)

            processed_paragraphs = 0



            last_student_number = None

            last_student_name = None



            for para in doc.paragraphs:

                processed_paragraphs += 1

                if processed_paragraphs % 1000 == 0:  # Log every 1000 paragraphs

                    progress = 10 + (processed_paragraphs / total_paragraphs * 60)  # Scale from 10% to 70%

                    logger.info(f"Processed {processed_paragraphs:,} paragraphs ({(processed_paragraphs/total_paragraphs*100):.1f}%)")

                    self.update_progress(progress, f"Processing paragraph {processed_paragraphs:,} of {total_paragraphs:,}")



                text = para.text.strip()

                if not text:

                    continue



                logger.info(f"Checking text: {text[:200]}...")



                # Detect student number first

                student_name, student_number = None, None

                for num in student_dict:

                    if num in text:

                        student_name, student_number = student_dict[num], num

                        last_student_name = student_name  # Store last detected student

                        last_student_number = student_number

                        logger.info(f"Found student: {student_name} ({student_number})")

                        break



                # If no student number, use the last detected one

                if not student_number and last_student_number:

                    student_name, student_number = last_student_name, last_student_number

                    logger.info(f"Using last detected student: {student_name} ({student_number})")



                # If we still don't have a student number, SKIP processing

                if not student_number:

                    logger.warning(f"⚠ No student number found in text: {text[:100]}")

                    continue



                # Now process violations/resolutions

                is_violation = any(kw in text.lower() for kw in violation_keywords)

                is_resolution = any(kw in text.lower() for kw in resolution_keywords)



                if is_violation or is_resolution:

                    logger.info(f"Matched keywords: Violations {is_violation}, Resolutions {is_resolution}")



                    # Extract date

                    date_match = re.search(r'(?:on|dtd)\s+([A-Za-z]+\.?\s+\d{1,2},?\s*\d{4}|\d{1,2}[-/]\d{1,2}[-/]\d{2,4})', text)

                    entry_date = date_match.group(1) if date_match else ""

                    logger.info(f"Date found: {entry_date}" if entry_date else "⚠ No date found in text")



                    # Extract name using regex

                    name_match = re.search(r'(?i)([A-Za-z\s,.-]+),\s*([A-Za-z\s-]+?)(?:\s+([A-Z])\.?)(?:\s|$)', text)

                    last_name = ""

                    first_name = ""

                    middle_initial = ""



                    if name_match:

                        last_name = name_match.group(1).strip()

                        first_name = name_match.group(2).strip()

                        middle_initial = name_match.group(3).strip() if name_match.group(3) else ""

                        logger.info(f"Extracted name: {last_name}, {first_name} {middle_initial}")

                    else:

                        logger.warning("No valid name found in text")



                    # Extract college and course information

                    college = ""

                    course = ""



                    # Assuming college and course information is in the text, extract them

                    college_match = re.search(r'College of ([A-Za-z\s]+)', text)

                    course_match = re.search(r'Course: ([A-Za-z\s]+)', text)



                    if college_match:

                        college = college_match.group(1).strip()

                        logger.info(f"College found: {college}")

                    else:

                        logger.warning("No college found in text")



                    if course_match:

                        course = course_match.group(1).strip()

                        logger.info(f"Course found: {course}")

                    else:

                        logger.warning("No course found in text")



                    row_data = [

                        student_number,

                        last_name,

                        first_name,

                        middle_initial,  # M.I.

                        college,  # College

                        course,  # Course

                        entry_date if is_violation else "",

                        text if is_violation else "",

                        entry_date if is_resolution else "",

                        text if is_resolution else "",

                        "",  # Resolver

                        text  # Remarks

                    ]



                    # Check for duplicates before appending

                    if not any(sheet.cell(row=r, column=1).value == student_number for r in range(2, last_row + 1)):

                        data_to_write.append(row_data)

                        logger.info(f"Successfully added entry for student {student_number}")

                    else:

                        logger.info(f"Duplicate entry found for student {student_number}, skipping...")



            # Log sample of prepared data before writing

            logger.info(f"Paragraph processing complete. Found {len(data_to_write)} entries to write.")

            if data_to_write:

                logger.info("Sample of first 5 entries to write:")

                for idx, entry in enumerate(data_to_write[:5]):

                    logger.info(f"Entry {idx + 1}: {entry}")

            else:

                logger.warning("No data was prepared for writing!")



            # Additional verification before Excel writing

            logger.info(f"Final verification - Total rows prepared: {len(data_to_write)}")

            if len(data_to_write) > 0:

                logger.info(f"Final data sample: {data_to_write[:2]}")



            self.update_progress(70, "Writing data to Excel file...")



            # Writing to Excel with enhanced logging

            total_rows = len(data_to_write)

            logger.info(f"Beginning Excel write process. Total rows to write: {total_rows}")

            logger.info("Attempting intermediate save before writing data...")



            # Force an intermediate save before writing

            try:

                logger.info("Attempting intermediate save before writing data...")

                workbook.save(self.excel_file)

                logger.info("Intermediate save completed successfully.")

            except Exception as e:

                logger.error(f"Error during intermediate save: {str(e)}")

                raise



            if not data_to_write:

                logger.warning(" No data to write to Excel!")

                return  # Stop early if there's no data



            for current_row, row_data in enumerate(data_to_write, start=last_row + 1):

                logger.info(f"Writing row {current_row}: {row_data}")



                if current_row % 100 == 0:  # Update progress every 100 rows

                    progress = 70 + (current_row / total_rows * 20)  # Scale from 70% to 90%

                    logger.info(f"Written {current_row:,} rows ({(current_row/total_rows*100):.1f}%)")

                    self.update_progress(progress, f"Writing row {current_row:,} of {total_rows:,}")



                    # Force intermediate save every 100 rows

                    try:

                        logger.info(f"Attempting intermediate save at row {current_row}...")

                        workbook.save(self.excel_file)

                        logger.info("Intermediate save completed successfully.")

                    except Exception as e:

                        logger.error(f"Error during intermediate save at row {current_row}: {str(e)}")

                        raise



                # Write each column of row_data to the Excel sheet

                for col_index, value in enumerate(row_data, start=1):

                    logger.debug(f"Writing to cell ({current_row}, {col_index}): {value}")



                    cell = sheet.cell(row=current_row, column=col_index)



                    # Check if the cell is part of a merged range

                    for merged_range in sheet.merged_cells.ranges:

                        if cell.coordinate in merged_range:

                            # Only write to the top-left cell of the merged range

                            top_left_cell = sheet.cell(row=merged_range.min_row, column=merged_range.min_col)

                            if top_left_cell.coordinate == cell.coordinate:

                                top_left_cell.value = value

                                top_left_cell.alignment = openpyxl.styles.Alignment(wrap_text=True)

                            break  # Stop checking merged ranges

                        else:  # Normal cells (not merged)

                            cell.value = value

                            cell.alignment = openpyxl.styles.Alignment(wrap_text=True)



            # Autofit column widths

            logger.info("Adjusting column widths...")

            self.update_progress(90, "Adjusting column widths...")



            for column in sheet.columns:

                max_length = 0

                column_letter = get_column_letter(column[0].column)

                for cell in column:

                    try:

                        if cell.value:

                            max_length = max(max_length, len(str(cell.value)))

                    except Exception as e:

                        logger.error(f"Error calculating max length for column {column_letter}: {str(e)}")

                adjusted_width = min(max_length + 2, 50)

                sheet.column_dimensions[column_letter].width = adjusted_width



            # Save the workbook with enhanced error logging

            logger.info("Attempting final save of Excel file...")

            self.update_progress(95, "Saving Excel file...")



            try:

                workbook.save(self.excel_file)

                logger.info("Excel file saved successfully!")

                messagebox.showinfo("Success", "Conversion completed successfully!")

            except Exception as e:

                logger.error(f"Error saving Excel file: {str(e)}")

                messagebox.showerror("Error", f"Failed to save Excel file: {str(e)}")



        except Exception as e:

            logger.error(f"Error during document conversion: {str(e)}", exc_info=True)

            messagebox.showerror("Error", f"An error occurred during document conversion:\n{str(e)}")

        finally:

            self.is_converting = False

            self.progress_var.set(0)

            self.progress_label.config(text="Ready to convert")

            self.root.update()



    def compare_documents(self):

        """Compare converted Excel with original DOCX/Excel for verification"""

        try:

            # First, let user select the Excel file

            excel_file = filedialog.askopenfilename(

                title="Select First Excel File",

                filetypes=[("Excel Files", "*.xlsx")]

            )

            if not excel_file:

                return

                

            # Then select the file to compare against (Word or Excel)

            source_file = filedialog.askopenfilename(

                title="Select File to Compare Against",

                filetypes=[("Word and Excel Files", "*.docx *.xlsx"), ("Word Documents", "*.docx"), ("Excel Files", "*.xlsx")]

            )

            if not source_file:

                return

            

            self.progress_label.config(text="Starting document comparison...")

            self.root.update()

            time.sleep(1)

            

            try:

                # Read first Excel file with explicit data types

                df = pd.read_excel(excel_file, dtype={'Student Number': str})

                # Convert any NaN values to empty strings

                df = df.fillna("")

            except Exception as e:

                logger.error(f"Error reading Excel file: {str(e)}")

                messagebox.showerror("Error", "Failed to read Excel file. Please ensure it's in the correct format.")

                return

            

            # Determine file type and read accordingly

            file_extension = os.path.splitext(source_file)[1].lower()

            

            try:

                if file_extension == '.docx':

                # Read DOCX

                    doc = Document(source_file)

                    # Process DOCX content

                    current_text = []

                    for para in doc.paragraphs:

                        text = para.text.strip()

                        if text:

                            current_text.append(text)

                    source_content = ' '.join(current_text)

                    is_docx = True

                else:

                    # Read second Excel file

                    source_df = pd.read_excel(source_file, dtype={'Student Number': str})

                    source_df = source_df.fillna("")

                    # Convert all student numbers to string and extract only digits

                    source_df['Student Number'] = source_df['Student Number'].astype(str).apply(lambda x: ''.join(filter(str.isdigit, x)))

                    is_docx = False

                    

            except Exception as e:

                logger.error(f"Error reading comparison file: {str(e)}")

                messagebox.showerror("Error", f"Failed to read comparison file. Please ensure it's a valid {'Word' if file_extension == '.docx' else 'Excel'} file.")

                return

            

            # Initialize comparison results

            comparison_results = {

                "total_records": len(df),

                "matched_records": 0,

                "mismatched_records": [],

                "missing_records": [],

                "extra_records": []

            }

            

            # Create a temporary text file for logging comparison results

            log_file = os.path.splitext(excel_file)[0] + '_comparison.txt'

            

            with open(log_file, 'w', encoding='utf-8') as f:

                f.write("Document Comparison Report\n")

                f.write("=" * 50 + "\n\n")

                f.write(f"First Excel File: {os.path.basename(excel_file)}\n")

                f.write(f"Comparison File: {os.path.basename(source_file)}\n")

                f.write("=" * 50 + "\n\n")

                

                # Compare each record in Excel with source content

                for idx, row in df.iterrows():

                    self.progress_var.set((idx + 1) / len(df) * 100)

                    self.progress_label.config(text=f"Comparing record {idx + 1} of {len(df)}...")

                    self.root.update()

                    

                    try:

                        # Student number pattern (just the raw number)

                        student_num = str(row["Student Number"]).strip()

                        if student_num and student_num != "nan":

                            # Remove any non-digit characters to get clean number

                            student_num = ''.join(filter(str.isdigit, student_num))

                            if student_num:

                                if is_docx:

                                    # Look for the exact student number without any labels in DOCX

                                    pattern = r'\b' + re.escape(student_num) + r'\b'

                                    match = re.search(pattern, source_content)

                                else:

                                    # Look for the student number in the second Excel file

                                    match = student_num in source_df['Student Number'].values

                                

                                if match:

                                    comparison_results["matched_records"] += 1

                                    f.write(f"Record {idx + 1}: Match found\n")

                                    f.write(f"  Student Number: {student_num}\n")

                                    

                                    if is_docx:

                                        # Get context around the match for verification (DOCX only)

                                        start = max(0, match.start() - 100)

                                        end = min(len(source_content), match.end() + 100)

                                        context = source_content[start:end].strip()

                                        f.write(f"  Context: ...{context}...\n")

                                    else:

                                        # For Excel comparison, show the matching row

                                        matching_row = source_df[source_df['Student Number'] == student_num].iloc[0]

                                        f.write("  Matching Record Details:\n")

                                        for col in matching_row.index:

                                            if matching_row[col]:

                                                f.write(f"    {col}: {matching_row[col]}\n")

                                    f.write("-" * 40 + "\n")

                                else:

                                    comparison_results["mismatched_records"].append(idx + 1)

                                    f.write(f"Record {idx + 1}: No match found\n")

                                    f.write(f"  Student Number: {student_num}\n")

                                    f.write(f"  WARNING: This student number was not found in the {'Word document' if is_docx else 'comparison Excel file'}\n")

                                    f.write("-" * 40 + "\n")

                            else:

                                comparison_results["mismatched_records"].append(idx + 1)

                                f.write(f"Record {idx + 1}: Invalid student number\n")

                                f.write(f"  Raw input: {row['Student Number']}\n")

                                f.write("  WARNING: Student number contains no digits\n")

                                f.write("-" * 40 + "\n")

                        else:

                            comparison_results["mismatched_records"].append(idx + 1)

                            f.write(f"Record {idx + 1}: Missing student number\n")

                            f.write("  WARNING: No student number provided in Excel\n")

                            f.write("-" * 40 + "\n")

                            

                    except Exception as e:

                        logger.error(f"Error processing row {idx + 1}: {str(e)}")

                        comparison_results["mismatched_records"].append(idx + 1)

                        f.write(f"Record {idx + 1}: Error during comparison\n")

                        f.write(f"  Error: {str(e)}\n")

                        f.write("-" * 40 + "\n")

                        continue

                

                # Write summary

                f.write("\nComparison Summary\n")

                f.write("=" * 50 + "\n")

                f.write(f"Total Records in First Excel: {comparison_results['total_records']}\n")

                f.write(f"Matched Records: {comparison_results['matched_records']}\n")

                f.write(f"Mismatched Records: {len(comparison_results['mismatched_records'])}\n")

                

                if comparison_results["mismatched_records"]:

                    f.write("\nMismatched Record Numbers:\n")

                    f.write(", ".join(map(str, comparison_results["mismatched_records"])))

            

            # Show results

            if comparison_results["total_records"] > 0:

                match_percentage = (comparison_results["matched_records"] / comparison_results["total_records"]) * 100

            else:

                match_percentage = 0

            

            messagebox.showinfo("Comparison Complete", 

                f"Document comparison completed!\n\n"

                f"Total Records: {comparison_results['total_records']}\n"

                f"Matched Records: {comparison_results['matched_records']}\n"

                f"Match Percentage: {match_percentage:.1f}%\n\n"

                f"Detailed report saved to:\n{os.path.basename(log_file)}"

            )

            

            # Open the log file

            os.startfile(log_file)

            

        except Exception as e:

            logger.error(f"Error during comparison: {str(e)}", exc_info=True)

            messagebox.showerror("Error", f"An error occurred during comparison:\n{str(e)}")

        finally:

            self.progress_var.set(0)

            self.progress_label.config(text="Ready to convert")

            self.root.update()



    def clear_form(self):

        """Reset all form fields to their default values"""

        self.word_file = None

        self.excel_file = None

        self.word_label.config(text="No file selected")

        self.excel_label.config(text="Default: Same folder as Word file")

        self.progress_var.set(0)

        self.progress_label.config(text="Ready to convert")



    def on_closing(self):

        """Handle window close event"""

        if self.is_converting:

            if not messagebox.askokcancel("Quit", "Conversion is in progress. Are you sure you want to quit?"):

                return

        

        self.root.destroy()



    def parse_student_info(self, text: str) -> dict:

        """Extract student information including College and Course from text."""

        info = {

            "Student Number": "",

            "Last Name": "",

            "First Name": "",

            "M.I.": "",

            "College": "",

            "Course": ""

        }



        # Extract College from known abbreviations or full names

        college_match = re.search(

            r'\b(CA|CASBE|CBA|CED|CE|CHASS|CISTM|CL|CM|CN|CPT|CS|CTHM|GSL)\b', text, re.IGNORECASE)

        

        if college_match:

            info["College"] = college_match.group(1).upper()



        # Extract Course from known abbreviations

        course_match = re.search(

            r'\b(BS-ARCH|BAC|BS-PSY|BSBA-BE|BSBA-FM|BSBA-HRM|BSBA-MM|BSBA-OM|BSEd-Eng|BSEd-Math|BSEd-Sci|BSCE|BSCpE|BSCS|BSECE|BSEE|BSME)\b',

            text, re.IGNORECASE)



        if course_match:

            info["Course"] = course_match.group(1).upper()



        return info



    def update_excel(self):

        """Update existing Excel file with data from Word document with thorough checking"""

        try:

            # First, select the Excel file to update

            excel_file = filedialog.askopenfilename(

                title="Select Excel File to Update",

                filetypes=[("Excel Files", "*.xlsx")]

            )

            if not excel_file:

                return



            # Then select the Word document with additional information

            word_file = filedialog.askopenfilename(

                title="Select Word Document with Updates",

                filetypes=[("Word Documents", "*.docx")]

            )

            if not word_file:

                return



            self.progress_label.config(text="Loading documents into memory...")

            self.root.update()



            # Load entire Excel file into memory

            try:

                existing_df = pd.read_excel(excel_file, dtype={'Student Number': str})

                existing_df = existing_df.fillna("")

                existing_df['Student Number'] = existing_df['Student Number'].apply(

                    lambda x: ''.join(filter(str.isdigit, str(x)))

                )

                

                # Validate required columns exist

                required_columns = ['Student Number', 'Last Name', 'First Name', 'M.I.', 'College', 'Course']

                missing_columns = [col for col in required_columns if col not in existing_df.columns]

                if missing_columns:

                    raise ValueError(f"Missing required columns: {', '.join(missing_columns)}")

                

            except Exception as e:

                logger.error(f"Error reading Excel file: {str(e)}")

                messagebox.showerror("Error", "Failed to read Excel file. Please ensure it's in the correct format.")

                return



            # Load entire Word document into memory

            doc = Document(word_file)

            word_data = {}



            # Parse Word document and create a dictionary of student data

            for para in doc.paragraphs:

                text = para.text.strip()

                if not text:

                    continue  # Skip empty lines



                # Instead of extracting student info again, just process violations/resolutions

                for student_number, student_info in word_data.items():

                    if student_number in text:  # Check if this text relates to a student

                        # Extract violations and resolutions

                        is_violation = any(keyword in text.lower() for keyword in self.violation_keywords)

                        is_resolution = any(keyword in text.lower() for keyword in self.resolution_keywords)



                        # Extract student info

                        parsed_info = self.parse_student_info(text)



                        # Prepare row data for writing

                        row_data = [

                            student_number,  

                            parsed_info["Last Name"],  

                            parsed_info["First Name"],  

                            parsed_info["M.I."],  

                            parsed_info["College"],  # Insert College

                            parsed_info["Course"],   # Insert Course

                            entry_date if is_violation else "",  

                            text if is_violation else "",  

                            entry_date if is_resolution else "",  

                            text if is_resolution else "",  

                            "",  # Resolver

                            text  # Remarks

                        ]



                        # Append row data to the list to be written

                        data_to_write.append(row_data)



            # Update Excel file with found information

            self.progress_label.config(text="Updating Excel with found information...")

            self.root.update()



            updates_made = 0

            successful_updates = []

            still_incomplete = []

            

            # Create a log file for tracking updates

            log_file = os.path.splitext(excel_file)[0] + '_update_log.txt'

            with open(log_file, 'w', encoding='utf-8') as f:

                f.write(f"Update Log - {time.strftime('%Y-%m-%d %H:%M:%S')}\n")

                f.write("=" * 50 + "\n\n")



                for idx, row in existing_df.iterrows():

                    try:

                        student_num = row['Student Number']

                        if not student_num:  # Skip rows without student numbers

                            continue

                            

                        if student_num in word_data:

                            update_info = word_data[student_num]

                            row_updates = []

                            

                            # Update personal information fields if missing

                            for field in ['Last Name', 'First Name', 'M.I.', 'College', 'Course']:

                                current_value = str(row[field]).strip()

                                new_value = str(update_info.get(field, '')).strip()

                                

                                if not current_value and new_value:

                                    existing_df.at[idx, field] = new_value

                                    updates_made += 1

                                    row_updates.append(f"{field}: '{new_value}'")

                                    successful_updates.append({

                                        'Student Number': student_num,

                                        'Field': field,

                                        'New Value': new_value

                                    })

                            

                            # Log updates for this student

                            if row_updates:

                                f.write(f"Student {student_num}:\n")

                                for update in row_updates:

                                    f.write(f"  - Updated {update}\n")

                                f.write("-" * 40 + "\n")

                        else:

                            still_incomplete.append(student_num)

                            

                    except Exception as e:

                        logger.error(f"Error updating row {idx + 1}: {str(e)}")

                        f.write(f"Error processing student {student_num}: {str(e)}\n")

                        continue



                # Write summary to log

                f.write("\nUpdate Summary\n")

                f.write("=" * 50 + "\n")

                f.write(f"Total Updates Made: {updates_made}\n")

                f.write(f"Students Still Incomplete: {len(still_incomplete)}\n")

                if still_incomplete:

                    f.write("\nIncomplete Student Numbers:\n")

                    f.write(", ".join(map(str, still_incomplete)))



            # Save updated Excel file with error handling

            try:

                # Create backup with timestamp

                timestamp = time.strftime('%Y%m%d_%H%M%S')

                backup_file = excel_file.replace('.xlsx', f'_backup_{timestamp}.xlsx')

                existing_df.to_excel(backup_file, index=False)

                

                # Save updates

                existing_df.to_excel(excel_file, index=False)

                

                # Show completion message with more details

                message = (

                    f"Excel file has been updated successfully!\n\n"

                    f"Updates Made: {updates_made}\n"

                    f"Students Still Incomplete: {len(still_incomplete)}\n\n"

                    f"A backup has been saved as:\n"

                    f"{os.path.basename(backup_file)}\n\n"

                    f"A detailed log file has been created at:\n"

                    f"{os.path.basename(log_file)}"

                )

                messagebox.showinfo("Update Complete", message)

                

                # Open the log file

                os.startfile(log_file)

                

            except Exception as e:

                error_msg = f"Failed to save updates: {str(e)}\n\nPlease ensure the Excel file is not open in another program."

                logger.error(error_msg)

                messagebox.showerror("Error", error_msg)

                return



        except Exception as e:

            logger.error(f"Error during update: {str(e)}", exc_info=True)

            messagebox.showerror("Error", f"An error occurred during update:\n{str(e)}")

        finally:

            self.progress_var.set(0)

            self.progress_label.config(text="Ready to convert")

            self.root.update()



    def update_progress(self, progress_value, message=""):

        """Update progress bar and progress label"""

        self.progress_var.set(progress_value)

        self.progress_label.config(text=message)

        self.root.update()





if __name__ == "__main__":

    main()